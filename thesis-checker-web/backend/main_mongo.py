from fastapi import FastAPI, File, UploadFile, Depends, HTTPException, status, Form, BackgroundTasks
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from pydantic import BaseModel
from typing import Optional, List
from datetime import datetime, timedelta
from jose import JWTError, jwt
from passlib.context import CryptContext
from pymongo import MongoClient
from pymongo.errors import DuplicateKeyError
import os
import uuid
import json
import requests
from docx import Document
import PyPDF2
import markdown
from dotenv import load_dotenv

# 加载环境变量
load_dotenv()

# 配置
SECRET_KEY = os.getenv("SECRET_KEY", "your-secret-key-keep-it-safe-in-production")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 300
UPLOAD_DIR = "./uploads"
REPORT_DIR = "./reports"
REQUIREMENTS_DIR = "./requirements"
TEMPLATES_DIR = "./templates"

# MongoDB 配置
MONGODB_URL = os.getenv("MONGODB_URL", "mongodb://admin:password123@localhost:27017")
DATABASE_NAME = os.getenv("DATABASE_NAME", "thesis_checker")

# AI 配置
AI_API_KEY = os.getenv("AI_API_KEY", "")
AI_BASE_URL = os.getenv("AI_BASE_URL", "https://api.deepseek.com/v1")
AI_MODEL = os.getenv("AI_MODEL", "deepseek-chat")

# 创建目录
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(REPORT_DIR, exist_ok=True)
os.makedirs(REQUIREMENTS_DIR, exist_ok=True)
os.makedirs(TEMPLATES_DIR, exist_ok=True)

# MongoDB 连接
client = MongoClient(MONGODB_URL)
db = client[DATABASE_NAME]

# 创建索引
db.users.create_index("username", unique=True)
db.users.create_index("email", unique=True)
db.theses.create_index([("owner_id", 1), ("created_at", -1)])

# 密码加密
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="api/token")

app = FastAPI(title="毕业论文检查系统API", version="1.0.0")

# CORS配置
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic 模型
class Token(BaseModel):
    access_token: str
    token_type: str
    is_admin: bool = False
    username: str

class TokenData(BaseModel):
    username: Optional[str] = None

class UserBase(BaseModel):
    username: str
    email: str

class UserCreate(UserBase):
    password: str

class UserResponse(BaseModel):
    id: str
    username: str
    email: str
    is_admin: bool
    created_at: datetime

class ThesisResponse(BaseModel):
    id: str
    title: str
    filename: str
    file_type: str
    status: str
    created_at: datetime
    owner_id: str

class RequirementResponse(BaseModel):
    id: str
    name: str
    type: str
    major: Optional[str]
    created_at: datetime

class TemplateResponse(BaseModel):
    id: str
    name: str
    category: str
    description: Optional[str]
    filename: str
    file_type: str
    created_at: datetime

# 辅助函数
def get_password_hash(password: str):
    return pwd_context.hash(password)

def verify_password(plain_password: str, hashed_password: str):
    return pwd_context.verify(plain_password, hashed_password)

def get_user_by_username(username: str):
    return db.users.find_one({"username": username})

def get_user_by_email(email: str):
    return db.users.find_one({"email": email})

def get_user_by_id(user_id: str):
    return db.users.find_one({"_id": user_id})

def authenticate_user(username: str, password: str):
    user = get_user_by_username(username)
    if not user:
        return False
    if not verify_password(password, user["hashed_password"]):
        return False
    return user

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(token: str = Depends(oauth2_scheme)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
        token_data = TokenData(username=username)
    except JWTError:
        raise credentials_exception
    user = get_user_by_username(token_data.username)
    if user is None:
        raise credentials_exception
    return user

async def get_current_admin(current_user: dict = Depends(get_current_user)):
    if not current_user.get("is_admin"):
        raise HTTPException(status_code=403, detail="Not authorized as admin")
    return current_user

# 文件读取函数
def read_docx(file_path):
    doc = Document(file_path)
    content = []
    for i, para in enumerate(doc.paragraphs):
        content.append({
            'line': i + 1,
            'text': para.text,
            'style': para.style.name,
            'alignment': str(para.alignment) if para.alignment else None,
            'indent': para.paragraph_format.first_line_indent.pt if para.paragraph_format.first_line_indent else None,
            'line_spacing': para.paragraph_format.line_spacing
        })
    return content

def read_doc(file_path):
    """读取 .doc 旧格式 Word 文件"""
    try:
        return read_docx(file_path)
    except Exception:
        return []

def read_pdf(file_path):
    content = []
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            content.append({
                'page': page_num + 1,
                'text': text
            })
    return content

def read_md(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    content = []
    for i, line in enumerate(lines):
        text = line.strip()
        if text:
            content.append({
                'line': i + 1,
                'text': text,
                'style': 'Normal'
            })
    return content

def call_llm(prompt: str):
    """请求大模型接口"""
    if not AI_API_KEY:
        return {
            "summary": "AI 评价功能未配置 API_KEY",
            "content_suggestions": [
                {"description": "请检查摘要是否精炼", "suggestion": "摘要应包含研究目的、方法、结果和结论"}
            ]
        }
    
    try:
        headers = {
            "Authorization": f"Bearer {AI_API_KEY}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": AI_MODEL,
            "messages": [
                {"role": "system", "content": "你是一位资深的论文指导老师，请严格根据提供的规范要求对论文内容进行评价。"},
                {"role": "user", "content": prompt}
            ],
            "response_format": {"type": "json_object"}
        }
        response = requests.post(f"{AI_BASE_URL}/chat/completions", headers=headers, json=payload, timeout=120)
        res_data = response.json()
        if "choices" not in res_data:
            return {"summary": "AI 响应格式不正确", "ai_issues": []}
        content_str = res_data["choices"][0]["message"]["content"]
        return json.loads(content_str)
    except Exception as e:
        return {"summary": f"AI 服务暂时不可用: {str(e)}", "content_suggestions": []}

def check_thesis(content, requirements_text, templates_info=""):
    """深度论文检查"""
    issues = []
    
    # 基础格式检查
    for para in content:
        if isinstance(para, dict) and 'text' in para:
            text = para['text'].strip()
            if not text: continue
            
            # 标题长度
            if 'style' in para and para['style'].startswith('Heading'):
                if len(text) > 40:
                    issues.append({
                        'position': f"第{para['line']}行 ({para['style']})",
                        'type': '格式问题',
                        'description': f"标题过长（{len(text)}字）",
                        'suggestion': '建议标题精简至30字以内',
                        'severity': 'medium'
                    })
    
    # AI 分析
    sample_text = "\n".join([p['text'] for p in content[:100] if isinstance(p, dict)])
    
    prompt = f"""
    任务：请根据以下论文内容和格式规范进行深度审计。
    
    【格式规范】：{requirements_text}
    【论文内容采样】：{sample_text}
    
    请以 JSON 格式输出：{{
        "ai_issues": [
            {{"type": "格式/结构/内容", "position": "位置描述", "description": "问题描述", "suggestion": "修改建议", "severity": "high/medium/low"}}
        ],
        "summary": "总体评价总结"
    }}
    """
    
    ai_result = call_llm(prompt)
    
    if "ai_issues" in ai_result:
        for ai_issue in ai_result["ai_issues"]:
            issues.append(ai_issue)
    
    # 计算评分
    base_score = 100
    deduction = len([i for i in issues if i.get('severity') == 'high']) * 15
    deduction += len([i for i in issues if i.get('severity') == 'medium']) * 8
    deduction += len([i for i in issues if i.get('severity') == 'low']) * 3
    
    final_score = max(0, base_score - deduction)
    summary = ai_result.get("summary", f"自动检查完成，发现 {len(issues)} 个潜在问题。")

    return {
        'issues': issues,
        'total_issues': len(issues),
        'score': final_score,
        'summary': summary
    }

# 路由
@app.post("/api/token", response_model=Token)
async def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends()):
    user = authenticate_user(form_data.username, form_data.password)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user["username"]}, expires_delta=access_token_expires
    )
    return {
        "access_token": access_token,
        "token_type": "bearer",
        "is_admin": user.get("is_admin", False),
        "username": user["username"]
    }

@app.post("/api/auth/register", response_model=UserResponse)
def register_user(user: UserCreate):
    # 检查用户名是否已存在
    if get_user_by_username(user.username):
        raise HTTPException(status_code=400, detail="Username already registered")
    
    # 检查邮箱是否已存在
    if get_user_by_email(user.email):
        raise HTTPException(status_code=400, detail="Email already registered")
    
    hashed_password = get_password_hash(user.password)
    user_dict = {
        "username": user.username,
        "email": user.email,
        "hashed_password": hashed_password,
        "is_admin": False,
        "created_at": datetime.utcnow()
    }
    
    try:
        result = db.users.insert_one(user_dict)
        user_dict["_id"] = str(result.inserted_id)
        return user_dict
    except DuplicateKeyError:
        raise HTTPException(status_code=400, detail="Username or email already exists")

@app.get("/api/auth/profile", response_model=UserResponse)
def read_users_me(current_user: dict = Depends(get_current_user)):
    current_user["id"] = str(current_user["_id"])
    return current_user

# 论文相关接口
@app.post("/api/thesis/upload", response_model=ThesisResponse)
async def upload_thesis(
    background_tasks: BackgroundTasks,
    title: str = Form(...),
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    file_ext = os.path.splitext(file.filename)[1].lower()
    allowed_extensions = ['.docx', '.doc', '.pdf', '.md']
    if file_ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail="Unsupported file type")
    
    file_id = str(uuid.uuid4())
    filename = f"{file_id}{file_ext}"
    file_path = os.path.join(UPLOAD_DIR, filename)
    
    with open(file_path, "wb") as buffer:
        content = await file.read()
        buffer.write(content)
    
    thesis_dict = {
        "title": title,
        "filename": file.filename,
        "file_path": file_path,
        "file_type": file_ext[1:],
        "owner_id": str(current_user["_id"]),
        "status": "checking",
        "created_at": datetime.utcnow()
    }
    
    result = db.theses.insert_one(thesis_dict)
    thesis_dict["_id"] = str(result.inserted_id)
    
    # 启动后台检查任务
    background_tasks.add_task(perform_check_logic, str(result.inserted_id))
    
    return thesis_dict

def perform_check_logic(thesis_id: str):
    """异步执行论文检查逻辑"""
    try:
        thesis = db.theses.find_one({"_id": thesis_id})
        if not thesis:
            return

        db.theses.update_one({"_id": thesis_id}, {"$set": {"status": "checking"}})

        # 读取文件内容
        content = []
        if thesis["file_type"] == 'docx':
            content = read_docx(thesis["file_path"])
        elif thesis["file_type"] == 'doc':
            content = read_doc(thesis["file_path"])
        elif thesis["file_type"] == 'pdf':
            content = read_pdf(thesis["file_path"])
        elif thesis["file_type"] == 'md':
            content = read_md(thesis["file_path"])
        
        # 加载规范要求
        requirements = list(db.requirements.find())
        req_content = "\n".join([f"[{r['type']}] {r['name']}: {r.get('content', '')}" for r in requirements])
        
        # 执行检查
        if not content:
            result = {
                'issues': [{
                    'position': '全文',
                    'type': '内容缺失',
                    'description': '未能从上传的文件中提取到有效文本内容',
                    'suggestion': '请尝试另存为 .docx 格式后重新上传',
                    'severity': 'high'
                }],
                'total_issues': 1,
                'score': 0,
                'summary': '文件读取失败，无法进行 AI 审计'
            }
        else:
            result = check_thesis(content, req_content)
        
        # 保存结果
        db.theses.update_one(
            {"_id": thesis_id},
            {"$set": {
                "check_result": result,
                "status": "completed"
            }}
        )

        # 生成报告文件
        report_filename = f"report_{thesis_id}.md"
        report_path = os.path.join(REPORT_DIR, report_filename)
        
        report_content = f"# 毕业论文检查报告\n\n"
        report_content += f"## 基本信息\n"
        report_content += f"- 论文标题：{thesis['title']}\n"
        report_content += f"- 文件名：{thesis['filename']}\n"
        report_content += f"- 检查时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        report_content += f"- 综合评分：{result['score']}分\n"
        report_content += f"- 问题总数：{result['total_issues']}个\n\n"
        
        report_content += f"## 问题详情\n"
        for i, issue in enumerate(result['issues'], 1):
            report_content += f"### {i}. {issue['position']}\n"
            report_content += f"**类型**：{issue['type']}\n"
            report_content += f"**问题描述**：{issue['description']}\n"
            report_content += f"**修改建议**：{issue['suggestion']}\n\n"
        
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write(report_content)
        
        db.theses.update_one(
            {"_id": thesis_id},
            {"$set": {"report_path": report_path}}
        )
    except Exception as e:
        print(f"Background check failed for thesis {thesis_id}: {str(e)}")
        db.theses.update_one({"_id": thesis_id}, {"$set": {"status": "failed"}})

@app.post("/api/thesis/check/{thesis_id}")
def check_thesis_endpoint(
    thesis_id: str,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    thesis = db.theses.find_one({"_id": thesis_id, "owner_id": str(current_user["_id"])})
    if not thesis:
        raise HTTPException(status_code=404, detail="Thesis not found")
    
    db.theses.update_one({"_id": thesis_id}, {"$set": {"status": "checking"}})
    background_tasks.add_task(perform_check_logic, thesis_id)
    
    return {"message": "论文检查已启动", "thesis_id": thesis_id, "status": "checking"}

@app.get("/api/thesis/history", response_model=List[ThesisResponse])
def get_thesis_history(current_user: dict = Depends(get_current_user)):
    theses = list(db.theses.find({"owner_id": str(current_user["_id"])}).sort("created_at", -1))
    for thesis in theses:
        thesis["id"] = str(thesis["_id"])
    return theses

@app.get("/api/thesis/report/{thesis_id}")
def get_thesis_report(
    thesis_id: str,
    current_user: dict = Depends(get_current_user)
):
    thesis = db.theses.find_one({"_id": thesis_id, "owner_id": str(current_user["_id"])})
    if not thesis:
        raise HTTPException(status_code=404, detail="Thesis not found")
    
    if thesis.get("status") != "completed":
        raise HTTPException(status_code=400, detail="Check not completed")
    
    if not thesis.get("check_result"):
        raise HTTPException(status_code=404, detail="Report not found")
    
    return thesis["check_result"]

@app.get("/api/thesis/report/{thesis_id}/download")
def download_thesis_report(
    thesis_id: str,
    current_user: dict = Depends(get_current_user)
):
    thesis = db.theses.find_one({"_id": thesis_id, "owner_id": str(current_user["_id"])})
    if not thesis:
        raise HTTPException(status_code=404, detail="Thesis not found")
    
    report_path = thesis.get("report_path")
    if not report_path or not os.path.exists(report_path):
        raise HTTPException(status_code=404, detail="Report file not found")
    
    return FileResponse(
        path=report_path,
        filename=f"论文检查报告_{thesis['title']}.md",
        media_type="text/markdown"
    )

@app.get("/api/thesis/{thesis_id}", response_model=ThesisResponse)
def get_thesis_info(
    thesis_id: str,
    current_user: dict = Depends(get_current_user)
):
    thesis = db.theses.find_one({"_id": thesis_id, "owner_id": str(current_user["_id"])})
    if not thesis:
        raise HTTPException(status_code=404, detail="Thesis not found")
    thesis["id"] = str(thesis["_id"])
    return thesis

@app.delete("/api/thesis/{thesis_id}")
def delete_thesis(
    thesis_id: str,
    current_user: dict = Depends(get_current_user)
):
    thesis = db.theses.find_one({"_id": thesis_id, "owner_id": str(current_user["_id"])})
    if not thesis:
        raise HTTPException(status_code=404, detail="Thesis not found")
    
    # 删除本地文件
    try:
        if thesis.get("file_path") and os.path.exists(thesis["file_path"]):
            os.remove(thesis["file_path"])
        if thesis.get("report_path") and os.path.exists(thesis["report_path"]):
            os.remove(thesis["report_path"])
    except Exception as e:
        print(f"Error deleting files: {e}")
    
    db.theses.delete_one({"_id": thesis_id})
    return {"message": "Thesis deleted successfully"}

# 管理员接口
@app.post("/api/admin/requirements", response_model=RequirementResponse)
async def upload_requirement(
    name: str = Form(...),
    type: str = Form(...),
    major: Optional[str] = Form(None),
    file: UploadFile = File(...),
    current_admin: dict = Depends(get_current_admin)
):
    if type not in ['school', 'major']:
        raise HTTPException(status_code=400, detail="Invalid requirement type")
    
    file_ext = os.path.splitext(file.filename)[1].lower()
    allowed_extensions = ['.md', '.txt', '.docx', '.doc', '.pdf']
    if file_ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail="Unsupported file type")
    
    file_id = str(uuid.uuid4())
    filename = f"{file_id}{file_ext}"
    file_path = os.path.join(REQUIREMENTS_DIR, filename)
    
    with open(file_path, "wb") as buffer:
        content = await file.read()
        buffer.write(content)
    
    # 提取文本内容
    content_text = ""
    if file_ext == '.docx':
        doc = Document(file_path)
        content_text = "\n".join([para.text for para in doc.paragraphs])
    elif file_ext == '.pdf':
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            content_text = "\n".join([page.extract_text() or '' for page in reader.pages])
    else:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content_text = f.read()
    
    req_dict = {
        "name": name,
        "type": type,
        "major": major,
        "file_path": file_path,
        "content": content_text,
        "created_at": datetime.utcnow()
    }
    
    result = db.requirements.insert_one(req_dict)
    req_dict["_id"] = str(result.inserted_id)
    return req_dict

@app.get("/api/admin/requirements", response_model=List[RequirementResponse])
def get_requirements(current_admin: dict = Depends(get_current_admin)):
    requirements = list(db.requirements.find().sort("created_at", -1))
    for req in requirements:
        req["id"] = str(req["_id"])
    return requirements

@app.delete("/api/admin/requirements/{req_id}")
def delete_requirement(
    req_id: str,
    current_admin: dict = Depends(get_current_admin)
):
    req = db.requirements.find_one({"_id": req_id})
    if not req:
        raise HTTPException(status_code=404, detail="Requirement not found")
    
    if req.get("file_path") and os.path.exists(req["file_path"]):
        os.remove(req["file_path"])
    
    db.requirements.delete_one({"_id": req_id})
    return {"message": "Requirement deleted successfully"}

@app.get("/api/admin/users", response_model=List[UserResponse])
def get_admin_users(current_admin: dict = Depends(get_current_admin)):
    users = list(db.users.find().sort("created_at", -1))
    for user in users:
        user["id"] = str(user["_id"])
    return users

# 模板接口
TEMPLATE_CATEGORIES = ['full', 'cover', 'body', 'reference']

@app.post("/api/admin/templates", response_model=TemplateResponse)
async def upload_template(
    name: str = Form(...),
    category: str = Form(...),
    description: Optional[str] = Form(None),
    file: UploadFile = File(...),
    current_admin: dict = Depends(get_current_admin)
):
    if category not in TEMPLATE_CATEGORIES:
        raise HTTPException(status_code=400, detail=f"Invalid category. Allowed: {TEMPLATE_CATEGORIES}")

    file_ext = os.path.splitext(file.filename)[1].lower()
    allowed_extensions = ['.docx', '.doc', '.pdf', '.md', '.txt']
    if file_ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_ext}")

    file_id = str(uuid.uuid4())
    stored_filename = f"{file_id}{file_ext}"
    file_path = os.path.join(TEMPLATES_DIR, stored_filename)

    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())

    tpl_dict = {
        "name": name,
        "category": category,
        "description": description,
        "filename": file.filename,
        "file_path": file_path,
        "file_type": file_ext[1:],
        "created_at": datetime.utcnow()
    }

    result = db.templates.insert_one(tpl_dict)
    tpl_dict["_id"] = str(result.inserted_id)
    return tpl_dict

@app.get("/api/admin/templates", response_model=List[TemplateResponse])
def get_templates():
    templates = list(db.templates.find().sort("created_at", -1))
    for tpl in templates:
        tpl["id"] = str(tpl["_id"])
    return templates

@app.delete("/api/admin/templates/{tpl_id}")
def delete_template(
    tpl_id: str,
    current_admin: dict = Depends(get_current_admin)
):
    tpl = db.templates.find_one({"_id": tpl_id})
    if not tpl:
        raise HTTPException(status_code=404, detail="Template not found")
    if tpl.get("file_path") and os.path.exists(tpl["file_path"]):
        os.remove(tpl["file_path"])
    db.templates.delete_one({"_id": tpl_id})
    return {"message": "Template deleted successfully"}

@app.get("/api/templates/{tpl_id}/download")
def download_template(
    tpl_id: str
):
    tpl = db.templates.find_one({"_id": tpl_id})
    if not tpl:
        raise HTTPException(status_code=404, detail="Template not found")
    if not tpl.get("file_path") or not os.path.exists(tpl["file_path"]):
        raise HTTPException(status_code=404, detail="File not found on disk")
    return FileResponse(
        path=tpl["file_path"],
        filename=tpl["filename"],
        media_type="application/octet-stream"
    )

# 创建默认管理员账号
def create_default_admin():
    admin = get_user_by_username("admin")
    if not admin:
        hashed_password = get_password_hash("admin123")
        admin_dict = {
            "username": "admin",
            "email": "admin@example.com",
            "hashed_password": hashed_password,
            "is_admin": True,
            "created_at": datetime.utcnow()
        }
        try:
            db.users.insert_one(admin_dict)
            print("Default admin user created: admin/admin123")
        except DuplicateKeyError:
            pass

# 启动时创建默认管理员
create_default_admin()

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
