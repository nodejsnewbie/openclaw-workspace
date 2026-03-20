from fastapi import FastAPI, File, UploadFile, Depends, HTTPException, status, Form, BackgroundTasks
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer
from pydantic import BaseModel
from typing import Optional, List
from datetime import datetime, timedelta
from jose import JWTError, jwt
from passlib.context import CryptContext
import os
import uuid
import json
import requests
from docx import Document
import PyPDF2
import markdown
from dotenv import load_dotenv
from bson import ObjectId
from database import database, USERS_COLLECTION, THESES_COLLECTION, REQUIREMENTS_COLLECTION, TEMPLATES_COLLECTION, create_indexes

# 加载环境变量
load_dotenv()

# 配置
SECRET_KEY = os.getenv("SECRET_KEY", "your-secret-key-keep-it-safe-in-production")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 300
UPLOAD_DIR = "./uploads"
REPORT_DIR = "./reports"
REQUIREMENTS_DIR = "./requirements"
TEMPLATES_DIR = "./templates"

# AI 配置 (建议将 API_KEY 放入环境变量)
AI_API_KEY = os.getenv("AI_API_KEY", "")  # 用户需提供 API 密钥
AI_BASE_URL = os.getenv("AI_BASE_URL", "https://api.deepseek.com/v1") # 默认 DeepSeek，可配置
AI_MODEL = os.getenv("AI_MODEL", "deepseek-chat")

# 创建目录
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(REPORT_DIR, exist_ok=True)
os.makedirs(REQUIREMENTS_DIR, exist_ok=True)
os.makedirs(TEMPLATES_DIR, exist_ok=True)

# 密码加密
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="api/token")

app = FastAPI(title="毕业论文检查系统API", version="1.0.0")

# CORS配置
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic模型
class Token(BaseModel):
    access_token: str
    token_type: str
    is_admin: bool
    username: str

class TokenData(BaseModel):
    username: Optional[str] = None

class UserBase(BaseModel):
    username: str
    email: str

class UserCreate(UserBase):
    password: str

class UserResponse(UserBase):
    id: str
    is_admin: bool
    created_at: datetime
    class Config:
        from_attributes = True

class ThesisBase(BaseModel):
    title: str

class ThesisResponse(ThesisBase):
    id: str
    filename: str
    file_type: str
    status: str
    created_at: datetime
    class Config:
        from_attributes = True

class RequirementBase(BaseModel):
    name: str
    type: str
    major: Optional[str] = None

class RequirementResponse(RequirementBase):
    id: str
    created_at: datetime
    updated_at: datetime
    class Config:
        from_attributes = True

class TemplateResponse(BaseModel):
    id: str
    name: str
    category: str
    description: Optional[str]
    filename: str
    file_type: str
    created_at: datetime
    class Config:
        from_attributes = True

class CheckResult(BaseModel):
    issues: List[dict]
    total_issues: int
    score: int
    summary: str

# 工具函数
def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

async def get_user(username: str):
    user = await database[USERS_COLLECTION].find_one({"username": username})
    return user

async def authenticate_user(username: str, password: str):
    user = await get_user(username)
    if not user:
        return False
    if not verify_password(password, user["hashed_password"]):
        return False
    return user

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(token: str = Depends(oauth2_scheme)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
        token_data = TokenData(username=username)
    except JWTError:
        raise credentials_exception
    user = await get_user(username=token_data.username)
    if user is None:
        raise credentials_exception
    return user

async def get_current_admin(current_user: dict = Depends(get_current_user)):
    if not current_user.get("is_admin"):
        raise HTTPException(status_code=403, detail="Not authorized as admin")
    return current_user

def read_docx(file_path):
    doc = Document(file_path)
    content = []
    for i, para in enumerate(doc.paragraphs):
        content.append({
            'line': i + 1,
            'text': para.text,
            'style': para.style.name,
            'alignment': str(para.alignment) if para.alignment else None,
            'indent': para.paragraph_format.first_line_indent.pt if para.paragraph_format.first_line_indent else None,
            'line_spacing': para.paragraph_format.line_spacing
        })
    return content

def read_doc(file_path):
    """
    读取 .doc 旧格式 Word 文件。
    优先使用 Win32 COM 组件（需安装 Microsoft Word），
    COM 不可用时尝试用 python-docx 兼容读取，两者均失败则返回空内容。
    """
    # 方案A：通过 Word COM 读取（Windows 独有，效果最佳）
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize() # 初始化 COM
        abs_path = os.path.abspath(file_path)
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        doc = None
        try:
            doc = word.Documents.Open(abs_path, ReadOnly=True)
            paragraphs = []
            count = doc.Paragraphs.Count
            for i in range(1, count + 1):
                try:
                    para = doc.Paragraphs(i)
                    text = para.Range.Text.strip('\r\x07\n') # 清理文档特殊字符
                    if text:
                        paragraphs.append({
                            'line': i,
                            'text': text,
                            'style': str(para.Style.NameLocal),
                            'alignment': None,
                            'indent': None,
                            'line_spacing': None
                        })
                except Exception as pe:
                    print(f"Error reading paragraph {i}: {pe}")
            return paragraphs
        finally:
            if doc: doc.Close(False)
            word.Quit()
    except Exception as e:
        print(f"win32com read_doc error: {e}")
    finally:
        try:
            import pythoncom
            pythoncom.CoUninitialize()
        except: pass
    # 方案B：部分 .doc 文件实际是 Word 2003 XML 格式， python-docx 可尝试兼容读取
    try:
        return read_docx(file_path)
    except Exception:
        pass
    return []  # 无法提取时返回空列表，让检查逻辑正常运行

def read_pdf(file_path):
    content = []
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            content.append({
                'page': page_num + 1,
                'text': text
            })
    return content

def read_md(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    content = []
    for i, line in enumerate(lines):
        text = line.strip()
        if text:
            content.append({
                'line': i + 1,
                'text': text,
                'style': 'Normal' # Markdown 简单处理为正文风格
            })
    return content

def call_llm(prompt: str):
    """请求大模型接口"""
    if not AI_API_KEY:
        # 如果没有配置 API KEY，返回一个模拟的 AI 评价
        return {
            "summary": "AI 评价功能未配置 API_KEY。请在后台环境变量中设置 AI_API_KEY 开启深度评价。",
            "content_suggestions": [
                {"description": "请检查摘要是否精炼", "suggestion": "摘要应包含研究目的、方法、结果和结论"},
                {"description": "关键词选取是否准确", "suggestion": "关键词应能代表论文核心内容"}
            ]
        }
    
    try:
        headers = {
            "Authorization": f"Bearer {AI_API_KEY}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": AI_MODEL,
            "messages": [
                {"role": "system", "content": "你是一位资深的论文指导老师，请严格根据提供的规范要求和模板，对论文内容进行评价，并给出具体的改进建议。请返回 JSON 格式结果。"},
                {"role": "user", "content": prompt}
            ],
            "response_format": {"type": "json_object"}
        }
        response = requests.post(f"{AI_BASE_URL}/chat/completions", headers=headers, json=payload, timeout=120)
        res_data = response.json()
        if "choices" not in res_data:
            print(f"LLM Response Error: {res_data}")
            return {"summary": "AI 响应格式不正确", "ai_issues": []}
        content_str = res_data["choices"][0]["message"]["content"]
        return json.loads(content_str)
    except Exception as e:
        print(f"LLM Call Error: {str(e)}")
        return {"summary": f"AI 服务暂时不可用: {str(e)}", "content_suggestions": []}

def check_thesis(content, requirements_text, templates_info=""):
    """
    深度论文检查：
    1. 基础格式硬性检查 (代码逻辑)
    2. 模板框架一致性检查 (LLM)
    3. 内容质量与学术建议 (LLM)
    """
    issues = []
    
    # 1. 基础格式检查 (正则/逻辑)
    for para in content:
        if isinstance(para, dict) and 'text' in para:
            text = para['text'].strip()
            if not text: continue
                
            # 标题长度
            if 'style' in para and para['style'].startswith('Heading'):
                if len(text) > 40:
                    issues.append({
                        'position': f"第{para['line']}行 ({para['style']})",
                        'type': '格式问题',
                        'description': f"标题过长（{len(text)}字）",
                        'suggestion': '建议标题精简至30字以内',
                        'severity': 'medium'
                    })
            
            # 正文缩进检查
            if 'style' in para and (para['style'] == 'Normal' or para['style'] == 'Body Text'):
                if 'indent' in para and (para['indent'] is None or para['indent'] < 5):
                     issues.append({
                        'position': f"第{para['line']}行",
                        'type': '格式问题',
                        'description': '正文段落首行未缩进',
                        'suggestion': '请设置首行缩进 2 字符',
                        'severity': 'medium'
                    })

    # 2. 构造大模型 Prompt 进行深度分析
    # 只抽取部分文本样本交给大模型以节省 Token 并提高速度
    sample_text = "\n".join([p['text'] for p in content[:100] if isinstance(p, dict)]) # 取前100段
    
    prompt = f"""
    任务：请根据以下论文内容、格式规范和书写模板进行深度审计。
    
    【格式规范】：
    {requirements_text}
    
    【书写模板信息】：
    {templates_info}
    
    【论文内容采样】：
    {sample_text}
    
    请严格检查：
    1. 目录结构是否与书写模板一致？是否有缺失章节。
    2. 内容逻辑：摘要是否规范，结论是否呼应。
    3. 学术性：语言是否学术，排版是否凌乱。
    
    请以 JSON 格式输出：
    {{
        "total_score_deduction": 10,
        "ai_issues": [
            {{"type": "格式/结构/内容", "position": "位置描述", "description": "问题描述", "suggestion": "修改建议", "severity": "high/medium/low"}}
        ],
        "summary": "总体评价总结",
        "detailed_advice": "详细的改进指导"
    }}
    """
    
    ai_result = call_llm(prompt)
    
    # 合并 AI 发现的问题
    if "ai_issues" in ai_result:
        for ai_issue in ai_result["ai_issues"]:
            issues.append(ai_issue)
            
    # 计算综合评分
    base_score = 100
    deduction = len([i for i in issues if i.get('severity') == 'high']) * 15
    deduction += len([i for i in issues if i.get('severity') == 'medium']) * 8
    deduction += len([i for i in issues if i.get('severity') == 'low']) * 3
    
    final_score = max(0, base_score - deduction)
    
    summary = ai_result.get("summary", f"自动检查完成，发现 {len(issues)} 个潜在问题。")
    if "detailed_advice" in ai_result:
        summary += "\n\n详细改进意见：\n" + ai_result["detailed_advice"]

    return {
        'issues': issues,
        'total_issues': len(issues),
        'score': final_score,
        'summary': summary
    }

# 启动事件：创建索引
@app.on_event("startup")
async def startup_event():
    await create_indexes()
    # 创建默认管理员账号
    await create_default_admin()

# 路由
@app.post("/api/token", response_model=Token)
async def login_for_access_token(
    username: str = Form(...),
    password: str = Form(...),
    grant_type: str = Form(default="password"),
):
    user = await authenticate_user(username, password)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user["username"]}, expires_delta=access_token_expires
    )
    return {
        "access_token": access_token,
        "token_type": "bearer",
        "is_admin": user["is_admin"],
        "username": user["username"]
    }

@app.post("/api/auth/register", response_model=UserResponse)
async def register_user(user: UserCreate):
    db_user = await get_user(username=user.username)
    if db_user:
        raise HTTPException(status_code=400, detail="Username already registered")
    
    # 检查邮箱是否已存在
    existing_email = await database[USERS_COLLECTION].find_one({"email": user.email})
    if existing_email:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    hashed_password = get_password_hash(user.password)
    user_doc = {
        "username": user.username,
        "email": user.email,
        "hashed_password": hashed_password,
        "is_admin": False,
        "created_at": datetime.utcnow()
    }
    result = await database[USERS_COLLECTION].insert_one(user_doc)
    user_doc["id"] = str(result.inserted_id)
    return user_doc

@app.get("/api/auth/profile", response_model=UserResponse)
async def read_users_me(current_user: dict = Depends(get_current_user)):
    current_user["id"] = str(current_user["_id"])
    return current_user

# 论文相关接口
@app.post("/api/thesis/upload", response_model=ThesisResponse)
async def upload_thesis(
    background_tasks: BackgroundTasks,
    title: str = Form(...),
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
):
    file_ext = os.path.splitext(file.filename)[1].lower()
    allowed_extensions = ['.docx', '.doc', '.pdf', '.md']
    if file_ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail="Unsupported file type")
    
    file_id = str(uuid.uuid4())
    filename = f"{file_id}{file_ext}"
    file_path = os.path.join(UPLOAD_DIR, filename)
    
    with open(file_path, "wb") as buffer:
        content = await file.read()
        buffer.write(content)
    
    thesis_doc = {
        "title": title,
        "filename": file.filename,
        "file_path": file_path,
        "file_type": file_ext[1:],
        "owner_id": str(current_user["_id"]),
        "status": "checking"
    }
    result = await database[THESES_COLLECTION].insert_one(thesis_doc)
    thesis_doc["id"] = str(result.inserted_id)
    
    # 启动后台检查任务
    background_tasks.add_task(perform_check_logic, str(result.inserted_id))
    
    return thesis_doc

def perform_check_logic(thesis_id: str):
    """异步执行论文检查逻辑 (同步版本，用于后台任务)"""
    import asyncio
    asyncio.run(_perform_check_logic_async(thesis_id))

async def _perform_check_logic_async(thesis_id: str):
    """异步执行论文检查逻辑"""
    from bson import ObjectId
    
    thesis = await database[THESES_COLLECTION].find_one({"_id": ObjectId(thesis_id)})
    if not thesis:
        return

    await database[THESES_COLLECTION].update_one(
        {"_id": ObjectId(thesis_id)},
        {"$set": {"status": "checking"}}
    )

    # 读取文件内容
    content = []
    if thesis['file_type'] == 'docx':
        content = read_docx(thesis['file_path'])
    elif thesis['file_type'] == 'doc':
        content = read_doc(thesis['file_path'])
    elif thesis['file_type'] == 'pdf':
        content = read_pdf(thesis['file_path'])
    elif thesis['file_type'] == 'md':
        content = read_md(thesis['file_path'])
    
    # 加载规范要求
    requirements = await database[REQUIREMENTS_COLLECTION].find().to_list(length=None)
    req_content = "\n".join([f"[{r['type']}] {r['name']}: {r['content']}" for r in requirements])
    
    # 加载书写模板信息
    templates = await database[TEMPLATES_COLLECTION].find().to_list(length=None)
    templates_info = "\n".join([f"模板类别: {t['category']}, 名称: {t['name']}, 描述: {t.get('description', '')}" for t in templates])
    
    # 执行检查
    print(f"Starting check for thesis {thesis_id}, content paragraphs: {len(content)}")
    if not content:
        # 如果内容为空，构造一个带警告的虚拟结果
        result = {
            'issues': [{
                'position': '全文',
                'type': '内容缺失',
                'description': '未能从上传的文件中提取到有效文本内容。请确认文件不是扫描件或加密文档。',
                'suggestion': '请尝试另存为 .docx 格式后重新上传。',
                'severity': 'high'
            }],
            'total_issues': 1,
            'score': 0,
            'summary': '文件读取失败，无法进行 AI 审计。'
        }
    else:
        result = check_thesis(content, req_content, templates_info)
    
    # 保存结果
    await database[THESES_COLLECTION].update_one(
        {"_id": ObjectId(thesis_id)},
        {"$set": {
            "check_result": json.dumps(result, ensure_ascii=False),
            "status": "completed"
        }}
    )

    # 生成报告文件
    report_filename = f"report_{thesis_id}.md"
    report_path = os.path.join(REPORT_DIR, report_filename)
    
    report_content = f"# 毕业论文检查报告\n\n"
    report_content += f"## 基本信息\n"
    report_content += f"- 论文标题：{thesis['title']}\n"
    report_content += f"- 文件名：{thesis['filename']}\n"
    report_content += f"- 检查时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
    report_content += f"- 综合评分：{result['score']}分\n"
    report_content += f"- 问题总数：{result['total_issues']}个\n\n"
    
    report_content += f"## 问题详情\n"
    for i, issue in enumerate(result['issues'], 1):
        report_content += f"### {i}. {issue['position']}\n"
        report_content += f"**类型**：{issue['type']}\n"
        report_content += f"**问题描述**：{issue['description']}\n"
        report_content += f"**修改建议**：{issue['suggestion']}\n\n"
    
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(report_content)
    
    await database[THESES_COLLECTION].update_one(
        {"_id": ObjectId(thesis_id)},
        {"$set": {"report_path": report_path}}
    )

@app.post("/api/thesis/check/{thesis_id}")
async def check_thesis_endpoint(
    thesis_id: str,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user),
):
    thesis = await database[THESES_COLLECTION].find_one({
        "_id": ObjectId(thesis_id),
        "owner_id": str(current_user["_id"])
    })
    if not thesis:
        raise HTTPException(status_code=404, detail="Thesis not found")
    
    await database[THESES_COLLECTION].update_one(
        {"_id": ObjectId(thesis_id)},
        {"$set": {"status": "checking"}}
    )
    
    # 启动后台任务
    background_tasks.add_task(perform_check_logic, thesis_id)
    
    return {"message": "论文检查已启动，请稍后查看结果", "thesis_id": thesis_id, "status": "checking"}

@app.get("/api/thesis/history", response_model=List[ThesisResponse])
async def get_thesis_history(current_user: dict = Depends(get_current_user)):
    print(f"Fetching history for user: {current_user['username']}")
    theses = await database[THESES_COLLECTION].find({
        "owner_id": str(current_user["_id"])
    }).sort("created_at", -1).to_list(length=None)
    
    result = []
    for t in theses:
        t["id"] = str(t["_id"])
        result.append(t)
    return result

@app.get("/api/thesis/report/{thesis_id}")
async def get_thesis_report(
    thesis_id: str,
    current_user: dict = Depends(get_current_user),
):
    thesis = await database[THESES_COLLECTION].find_one({
        "_id": ObjectId(thesis_id),
        "owner_id": str(current_user["_id"])
    })
    if not thesis:
        raise HTTPException(status_code=404, detail="Thesis not found")
    
    if thesis.get("status") != "completed":
        raise HTTPException(status_code=400, detail="Check not completed")
    
    if not thesis.get("check_result"):
        raise HTTPException(status_code=404, detail="Report not found")
    
    return json.loads(thesis["check_result"])

@app.get("/api/thesis/report/{thesis_id}/download")
async def download_thesis_report(
    thesis_id: str,
    current_user: dict = Depends(get_current_user),
):
    """下载生成的 Markdown 格式检查报告"""
    thesis = await database[THESES_COLLECTION].find_one({
        "_id": ObjectId(thesis_id),
        "owner_id": str(current_user["_id"])
    })
    if not thesis:
        raise HTTPException(status_code=404, detail="Thesis not found")
    
    report_path = thesis.get("report_path")
    if not report_path or not os.path.exists(report_path):
        raise HTTPException(status_code=404, detail="Report file not found")
    
    return FileResponse(
        path=report_path,
        filename=f"论文检查报告_{thesis['title']}.md",
        media_type="text/markdown"
    )

@app.get("/api/thesis/{thesis_id}", response_model=ThesisResponse)
async def get_thesis_info(
    thesis_id: str,
    current_user: dict = Depends(get_current_user),
):
    thesis = await database[THESES_COLLECTION].find_one({
        "_id": ObjectId(thesis_id),
        "owner_id": str(current_user["_id"])
    })
    if not thesis:
        raise HTTPException(status_code=404, detail="Thesis not found")
    thesis["id"] = str(thesis["_id"])
    return thesis

@app.delete("/api/thesis/{thesis_id}")
async def delete_thesis(
    thesis_id: str,
    current_user: dict = Depends(get_current_user),
):
    """删除论文记录及其关联文件"""
    thesis = await database[THESES_COLLECTION].find_one({
        "_id": ObjectId(thesis_id),
        "owner_id": str(current_user["_id"])
    })
    if not thesis:
        raise HTTPException(status_code=404, detail="Thesis not found")
    
    # 删除本地文件
    try:
        if thesis.get("file_path") and os.path.exists(thesis["file_path"]):
            os.remove(thesis["file_path"])
        if thesis.get("report_path") and os.path.exists(thesis["report_path"]):
            os.remove(thesis["report_path"])
    except Exception as e:
        print(f"Error deleting files for thesis {thesis_id}: {e}")
    
    await database[THESES_COLLECTION].delete_one({"_id": ObjectId(thesis_id)})
    return {"message": "Thesis deleted successfully"}

# 管理员接口
@app.post("/api/admin/requirements", response_model=RequirementResponse)
async def upload_requirement(
    name: str = Form(...),
    type: str = Form(...),
    major: Optional[str] = Form(None),
    file: UploadFile = File(...),
    current_admin: dict = Depends(get_current_admin),
):
    if type not in ['school', 'major']:
        raise HTTPException(status_code=400, detail="Invalid requirement type")
    
    file_ext = os.path.splitext(file.filename)[1].lower()
    allowed_extensions = ['.md', '.txt', '.docx', '.doc', '.pdf']
    if file_ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail="Unsupported file type")
    
    file_id = str(uuid.uuid4())
    filename = f"{file_id}{file_ext}"
    file_path = os.path.join(REQUIREMENTS_DIR, filename)
    
    with open(file_path, "wb") as buffer:
        content = await file.read()
        buffer.write(content)
    
    # 提取文本内容
    content = ""
    if file_ext == '.docx':
        doc = Document(file_path)
        content = "\n".join([para.text for para in doc.paragraphs])
    elif file_ext == '.doc':
        paras = read_doc(file_path)
        content = "\n".join([p['text'] for p in paras])
    elif file_ext == '.pdf':
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            content = "\n".join(page.extract_text() or '' for page in reader.pages)
    else:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
    
    req_doc = {
        "name": name,
        "type": type,
        "major": major,
        "file_path": file_path,
        "content": content,
        "created_at": datetime.utcnow(),
        "updated_at": datetime.utcnow()
    }
    result = await database[REQUIREMENTS_COLLECTION].insert_one(req_doc)
    req_doc["id"] = str(result.inserted_id)
    return req_doc

@app.get("/api/admin/requirements", response_model=List[RequirementResponse])
async def get_requirements(current_admin: dict = Depends(get_current_admin)):
    requirements = await database[REQUIREMENTS_COLLECTION].find().sort("created_at", -1).to_list(length=None)
    result = []
    for r in requirements:
        r["id"] = str(r["_id"])
        result.append(r)
    return result

@app.delete("/api/admin/requirements/{req_id}")
async def delete_requirement(
    req_id: str,
    current_admin: dict = Depends(get_current_admin),
):
    req = await database[REQUIREMENTS_COLLECTION].find_one({"_id": ObjectId(req_id)})
    if not req:
        raise HTTPException(status_code=404, detail="Requirement not found")
    
    if os.path.exists(req["file_path"]):
        os.remove(req["file_path"])
    
    await database[REQUIREMENTS_COLLECTION].delete_one({"_id": ObjectId(req_id)})
    return {"message": "Requirement deleted successfully"}

@app.get("/api/admin/users", response_model=List[UserResponse])
async def get_admin_users(current_admin: dict = Depends(get_current_admin)):
    """管理员获取所有用户列表"""
    users = await database[USERS_COLLECTION].find().sort("created_at", -1).to_list(length=None)
    result = []
    for u in users:
        u["id"] = str(u["_id"])
        result.append(u)
    return result

# ───────── 书写模板接口 ─────────

TEMPLATE_CATEGORIES = ['full', 'cover', 'body', 'reference']

@app.post("/api/admin/templates", response_model=TemplateResponse)
async def upload_template(
    name: str = Form(...),
    category: str = Form(...),
    description: Optional[str] = Form(None),
    file: UploadFile = File(...),
    current_admin: dict = Depends(get_current_admin),
):
    if category not in TEMPLATE_CATEGORIES:
        raise HTTPException(status_code=400, detail=f"Invalid category. Allowed: {TEMPLATE_CATEGORIES}")

    file_ext = os.path.splitext(file.filename)[1].lower()
    allowed_extensions = ['.docx', '.doc', '.pdf', '.md', '.txt']
    if file_ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_ext}. Allowed: {allowed_extensions}")

    file_id = str(uuid.uuid4())
    stored_filename = f"{file_id}{file_ext}"
    file_path = os.path.join(TEMPLATES_DIR, stored_filename)

    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())

    tpl_doc = {
        "name": name,
        "category": category,
        "description": description,
        "filename": file.filename,
        "file_path": file_path,
        "file_type": file_ext[1:],
        "created_at": datetime.utcnow()
    }
    result = await database[TEMPLATES_COLLECTION].insert_one(tpl_doc)
    tpl_doc["id"] = str(result.inserted_id)
    return tpl_doc

@app.get("/api/admin/templates", response_model=List[TemplateResponse])
async def get_templates():
    """所有用户均可查看模板列表"""
    templates = await database[TEMPLATES_COLLECTION].find().sort("created_at", -1).to_list(length=None)
    result = []
    for t in templates:
        t["id"] = str(t["_id"])
        result.append(t)
    return result

@app.delete("/api/admin/templates/{tpl_id}")
async def delete_template(
    tpl_id: str,
    current_admin: dict = Depends(get_current_admin),
):
    tpl = await database[TEMPLATES_COLLECTION].find_one({"_id": ObjectId(tpl_id)})
    if not tpl:
        raise HTTPException(status_code=404, detail="Template not found")
    if os.path.exists(tpl["file_path"]):
        os.remove(tpl["file_path"])
    await database[TEMPLATES_COLLECTION].delete_one({"_id": ObjectId(tpl_id)})
    return {"message": "Template deleted successfully"}

@app.get("/api/templates/{tpl_id}/download")
async def download_template(tpl_id: str):
    """任何已登录用户均可下载模板文件"""
    tpl = await database[TEMPLATES_COLLECTION].find_one({"_id": ObjectId(tpl_id)})
    if not tpl:
        raise HTTPException(status_code=404, detail="Template not found")
    if not os.path.exists(tpl["file_path"]):
        raise HTTPException(status_code=404, detail="File not found on disk")
    return FileResponse(
        path=tpl["file_path"],
        filename=tpl["filename"],
        media_type="application/octet-stream"
    )

# 创建默认管理员账号
async def create_default_admin():
    admin = await get_user(username="admin")
    if not admin:
        hashed_password = get_password_hash("admin123")
        admin_doc = {
            "username": "admin",
            "email": "admin@example.com",
            "hashed_password": hashed_password,
            "is_admin": True,
            "created_at": datetime.utcnow()
        }
        await database[USERS_COLLECTION].insert_one(admin_doc)

# 挂载前端静态文件（必须在所有 API 路由之后）
frontend_dist_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "frontend", "dist")
if os.path.exists(frontend_dist_path):
    app.mount("/", StaticFiles(directory=frontend_dist_path, html=True), name="frontend")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
